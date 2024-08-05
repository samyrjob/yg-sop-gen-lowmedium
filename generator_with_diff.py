import streamlit as st
import io
import os
import openai as ai
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import spacy
import json
import random
import pickle
import asyncio
import aiohttp

import requests
from bs4 import BeautifulSoup

import sys
# import logging

# # Configure the logging
# logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')


# Your Diffbot API token


PICKLE_FILE = 'scraped_data.pkl'


load_dotenv()  # take environment variables from .env.
DIFFBOT_API_TOKEN = os.getenv("DIFFBOT_API_TOKEN")
DIFFBOT_API_URL = 'https://api.diffbot.com/v3/article'

ai.api_key = os.getenv("OPENAI_API_KEY")
#test
nlp = spacy.blank("en")  # Create a blank English model
#^ Load the SpaCy English model
nlp_e = spacy.load("en_core_web_sm")

async def generate_sop4(template_text, res_text,programme,university, university_response, programme_response, germany_response):
    
    
    cgpa_score = await retrieve_cgpa_score(res_text)
    
    print("from the terminal : ", cgpa_score)
    
    if cgpa_score != 'unknown' and float(cgpa_score) >= 7: 
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
             {"role": "user", "content": "Generate a statement of purpose (SOP) based on the provided details, following a specific structure and style."},

            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Here is the resume {res_text} from which you have to take information about the student and please make sure that you do not add things from your own and stick to this {res_text} througout the sop"},
            {"role": "user" , "content": f"Another important that you should keep in mind while making the sop is that the sop is a letter of motivation to study in this {university} , so you have to generate the sop which should contian strong motivational points for choosing {programme}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"University name: {university}"},

            {"role": "user", "content": f" For the Introduction's paragraph, you must be using content from the {res_text} and follow the same structure and catchy words from from the introduction paragraph of this template {template_text}"},


            {"role": "user", "content": f"""Fot the Academic and experience Paragraph too you will be using content from the {res_text}: Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention all his acadamic education scores from the {res_text}
             Mention all the acadamic education scores as they are prensent in {res_text}
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list.
            You will be mentioning things like his certificates and Ilets score if they are present in the {res_text}
            """},


            {"role": "user", "content": f""" For the paragraph about the programme use strong motivational points for choosing {programme} and add information from {programme_response} into this paragraph, 
                Show a STRONG motivation by evoking a storytelling from your past that led you want to study this programme.
                If the past education of yours is related to the programme you are applying then you can consider them adding in this paragraph too
                You have to choose MINIMUM two out of these points :
                Story from a family experience
                Story from a social phenomenon and you want to bring solutions
                Story from previous internship
                Story from the studied modules in your Academic cursus
                Story from one of your extra-curricular activites
                You will also be adding What are the benifits you will be getting from this {programme} and the {university} 
                You must also select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Add learning outcomes form the curriculum of this {university} and these learning will be divided into two parts which are tecnical learning and interpersonal learning
                Explain Why this programme can help you build your career and prepare you for your future too.
                **Important:** Please do not invent any information not found in the `programme_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},


            # {"role": "user", "content": f""" For the paragraph about the University, you have to show the reasons for choosing {university} and try to write this like you want to apply for it(first person tone). And add information from {university_response} into this paragraph :
            #                                  You can use the following points to show the reason for choosing the university :  
            #                                                 Exact ranking of the University and the source of ranking
            #                                                 Number of students
            #                                                 facilities, faculties , campus location
            #                                                 Precise you have relatives and friends there
            #                                                 Names of Research centers and renowned Professors linked to the programme
                                                            
            
            # In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
            #                                                                                                             names of companies in the city that accept people from {programme} to work in       
            # """},
{"role": "user", "content": f"""For the paragraph about the University, you need to craft a compelling and personalized statement about why you want to apply to {university}. Write in the first person, as if you are the applicant. Incorporate detailed information from {university_response} into this paragraph, addressing the following points:

1. **University Ranking and Source:**
   - Include the exact ranking of {university} and cite the source of this ranking.

2. **Student Body and Facilities:**
   - Mention the number of students at {university}.
   - Describe the facilities, faculty, and campus location in detail.

3. **Personal Connections:**
   - Note if you have any relatives or friends at {university} and how that influences your choice.

4. **Research Centers and Professors:**
   - Highlight any notable research centers and renowned professors associated with the program you are interested in.

5. **City Attractions:**
   - Explain why you want to study in the city where {university} is located. Include some cool spots or attractions in the city that interest you.

6. **Local Job Market:**
   - Provide examples of companies in the city that hire graduates from the {programme}.

Make sure to weave all these elements into a cohesive and engaging paragraph that reflects your genuine interest in both the university and the city. Be detailed, specific, and personal in your response.

**Important:** Please do not invent any information not found in the `university_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

"""},
            {"role": "user", "content": f"""For the paragraph which is about your future career perspectives and aspirations post-study, you should mention how the programme of {programme}
             will give you the opportunity to work in the futur job you want to do. You should mention names of professional positions linked to the programme you want to work in.
             You can also add that you aim to work in Germany"""},


            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": f"""For the paragraph about Germany, you must mention the following points and add information from {germany_response} into this paragraph :
            Point 1. Explain why you chose to apply in Germany
            Point 2. Your Intention to stay in Germany because the field education comparative is far better to others destinations
            Point 3. Talk about the Good exposure, diversity and German culture 
            **Important:** Please do not invent any information not found in the `germany_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},

            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},

            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": f"Be sure to replace the correct information from {res_text} including the name."},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned', 'intricacies', 'close-knit', 'aligns', 'hands-on', 'enhance', 'foster', 'emphasis'"},

        ]

        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    
    
    elif cgpa_score != 'unknown' and float(cgpa_score) < 7:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
           {"role": "user", "content": "Generate a statement of purpose (SOP) based on the provided details, following a specific structure and style."},

            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Here is the resume {res_text} from which you have to take information about the student and please make sure that you do not add things from your own and stick to this {res_text} througout the sop"},
            {"role": "user" , "content": f"Another important that you should keep in mind while making the sop is that the sop is a letter of motivation to study in this {university} , so you have to generate the sop which should contian strong motivational points for choosing {programme}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"University name: {university}"},

            {"role": "user", "content": f" For the Introduction's paragraph, you must be using content from the {res_text} and follow the same structure and catchy words from from the introduction paragraph of this template {template_text}"},


            {"role": "user", "content": f"""Fot the Academic and experience Paragraph too you will be using content from the {res_text}: Discuss the candidate's academic and professional background :
            IF CGPA score is less than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
             Mention all his acadamic education scores from the {res_text}
             Mention all the acadamic education scores as they are prensent in {res_text}        
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list.
            You will be mentioning things like his certificates and Ilets score if they are present in the {res_text}
            """},


            {"role": "user", "content": f""" For the paragraph about the programme use strong motivational points for choosing {programme}, 
                Show a STRONG motivation by evoking a storytelling from your past that led you want to study this programme.
                If the past education of yours is related to the programme you are applying then you can consider them adding in this paragraph too
                You have to choose MINIMUM two out of these points :
                Story from a family experience
                Story from a social phenomenon and you want to bring solutions
                Story from previous internship
                Story from the studied modules in your Academic cursus
                Story from one of your extra-curricular activites
                You will also be adding What are the benifits you will be getting from this {programme} and the {university} 
                You must also select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Add learning outcomes form the curriculum of this {university} and these learning will be divided into two parts which are tecnical learning and interpersonal learning
                Explain Why this programme can help you build your career and prepare you for your future too 
            """},


            {"role": "user", "content": f""" For the paragraph about the University, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of Research centers and renowned Professors linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                        names of companies in the city that accept people from {programme} to work in       
            """},

            {"role": "user", "content": f"""For the paragraph which is about your future career perspectives and aspirations post-study, you should mention how the programme of {programme}
             will give you the opportunity to work in the futur job you want to do. You should mention names of professional positions linked to the programme you want to work in.
             You can also add that you aim to work in Germany"""},


            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": """For the paragraph about Germany, you must mention the following points :
            Point 1. Explain why you chose to apply in Germany
            Point 2. Your Intention to stay in Germany because the field education comparative is far better to others destinations
            Point 3. Talk about the Good exposure, diversity and German culture 
            """},

            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},

            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": f"Be sure to replace the correct information from {res_text} including the name."},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned', 'intricacies', 'close-knit', 'aligns', 'hands-on', 'enhance', 'foster', 'emphasis'"},


        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out
    else:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
                       {"role": "user", "content": "Generate a statement of purpose (SOP) based on the provided details, following a specific structure and style."},

            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Here is the resume {res_text} from which you have to take information about the student and please make sure that you do not add things from your own and stick to this {res_text} througout the sop"},
            {"role": "user" , "content": f"Another important that you should keep in mind while making the sop is that the sop is a letter of motivation to study in this {university} , so you have to generate the sop which should contian strong motivational points for choosing {programme}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"University name: {university}"},

            {"role": "user", "content": f" For the Introduction's paragraph, you must be using content from the {res_text} and follow the same structure and catchy words from from the introduction paragraph of this template {template_text}"},


            {"role": "user", "content": f"""Fot the Academic and experience Paragraph too you will be using content from the {res_text}: Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
             Mention all his acadamic education scores from the {res_text}
             Mention all the acadamic education scores as they are prensent in {res_text}
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list.
            You will be mentioning things like his certificates and Ilets score if they are present in the {res_text}
            """},


            {"role": "user", "content": f""" For the paragraph about the programme use strong motivational points for choosing {programme}, 
                Show a STRONG motivation by evoking a storytelling from your past that led you want to study this programme.
                If the past education of yours is related to the programme you are applying then you can consider them adding in this paragraph too
                You have to choose MINIMUM two out of these points :
                Story from a family experience
                Story from a social phenomenon and you want to bring solutions
                Story from previous internship
                Story from the studied modules in your Academic cursus
                Story from one of your extra-curricular activites
                You will also be adding What are the benifits you will be getting from this {programme} and the {university} 
                You must also select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Add learning outcomes form the curriculum of this {university} and these learning will be divided into two parts which are tecnical learning and interpersonal learning
                Explain Why this programme can help you build your career and prepare you for your future too 
            """},


            {"role": "user", "content": f""" For the paragraph about the University, you should retrieve the following data :
                                                            Exact ranking of the University and the source of ranking
                                                            Number of students
                                                            facilities, faculties , campus location
                                                            Precise you have relatives and friends there
                                                            Names of Research centers and renowned Professors linked to the programme
            You should use ALL these data and invent for each one an element of motivation make you want to integrate the University
            In this paragraph you should also mention why you want to study in the city of the University , mention :  Some cool spots in the city you would like to see
                                                                                                                        names of companies in the city that accept people from {programme} to work in       
            """},

            {"role": "user", "content": f"""For the paragraph which is about your future career perspectives and aspirations post-study, you should mention how the programme of {programme}
             will give you the opportunity to work in the futur job you want to do. You should mention names of professional positions linked to the programme you want to work in.
             You can also add that you aim to work in Germany"""},


            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": """For the paragraph about Germany, you must mention the following points :
            Point 1. Explain why you chose to apply in Germany
            Point 2. Your Intention to stay in Germany because the field education comparative is far better to others destinations
            Point 3. Talk about the Good exposure, diversity and German culture 
            """},

            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},

            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": f"Be sure to replace the correct information from {res_text} including the name."},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned', 'intricacies', 'close-knit', 'aligns', 'hands-on', 'enhance', 'foster', 'emphasis'"},

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write(response_out)
        return response_out



def extract_urls(query):

    # Perform the search query
    # query = "Hochschule Schmalkalden • (University) • Schmalkalden Master in Finance"
    
    
    url = f'https://www.google.com/search?q={requests.utils.quote(query)}'

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')

        # Find the first search result link
        links = []
        results = soup.find_all('div', {'class': 'yuRUbf'})
        for result in results[:2]:
            links.append(result.find('a')['href'])
            print(result.find('a')['href'])
        else:
            print("No results found.")
    else:
        print("Failed to retrieve the website content.")
        

    # Using slicing to access the first two elements
    if len(links) >= 2:
        sliced_links = links[0:2]
        return sliced_links[0], sliced_links[1]
    # # Print each element individually
    # for i, link in enumerate(sliced_links):
    #     print(f"Link {i + 1}: {link}")



async def extract3urls(query):
    url = f'https://www.google.com/search?q={requests.utils.quote(query)}'

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }

    async with aiohttp.ClientSession() as session:
        async with session.get(url, headers=headers) as response:
            if response.status == 200:
                html = await response.text()
                soup = BeautifulSoup(html, 'html.parser')

                # Find the first search result links
                links = []
                results = soup.find_all('div', {'class': 'yuRUbf'})
                for result in results[:10]:
                    link = result.find('a')['href']
                    links.append(link)
                    # print(link)  #^ temporary  
                
                if len(links) >= 10:
                    # Slice the first 10 elements
                    sliced_links = links[:10]
                    
                    # Get 3 unique random indices within the range of 0 to 9
                    random_indices = random.sample(range(10), 3)
                    
                    # Retrieve the links at the random indices
                    first_link = sliced_links[random_indices[0]]
                    second_link = sliced_links[random_indices[1]]
                    third_link = sliced_links[random_indices[2]]
                    
                    #^ temporary 
                    # print(f'{query} :', '  First link:', first_link, 'Second link:', second_link, 'Third link:', third_link)
                    
                    return first_link, second_link, third_link
                else:
                    return None, None, None
            else:
                print("Failed to retrieve the website content.")
                return None, None, None



async def extract3urls_orderofresults(query):
    url = f'https://www.google.com/search?q={requests.utils.quote(query)}'

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }

    async with aiohttp.ClientSession() as session:
        async with session.get(url, headers=headers) as response:
            if response.status == 200:
                html = await response.text()
                soup = BeautifulSoup(html, 'html.parser')

                # Find the first search result links
                links = []
                results = soup.find_all('div', {'class': 'yuRUbf'})
                for result in results[:10]:
                    link = result.find('a')['href']
                    links.append(link)
                    # print(link)  #^ temporary
                
                if len(links) >= 10:
                    # Slice the first 10 elements
                    sliced_links = links[:3]
                    

                    # Retrieve the links at the random indices
                    first_link = sliced_links[0]
                    second_link = sliced_links[1]
                    third_link = sliced_links[2]
                    
                    #^ temporary : 
                    # print(f'{query} :', 'First link:', first_link, 'Second link:', second_link, 'Third link:', third_link)
                    
                    return first_link, second_link, third_link
                else:
                    return None, None, None
            else:
                print("Failed to retrieve the website content.")
                return None, None, None


async def extract_1_url_from_urls(query):
    url = f'https://www.google.com/search?q={requests.utils.quote(query)}'

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }

    async with aiohttp.ClientSession() as session:
        async with session.get(url, headers=headers) as response:
            if response.status == 200:
                html = await response.text()
                soup = BeautifulSoup(html, 'html.parser')

                # Find the first search result links
                links = []
                results = soup.find_all('div', {'class': 'yuRUbf'})
                for result in results[:5]:
                    link = result.find('a')['href']
                    links.append(link)
                    # print(link) #^ temporary
                
                if len(links) >= 5:
                    # Slice the first 10 elements
                    sliced_links = links[:5]
                    
                    # Get 3 unique random indices within the range of 0 to 4
                 
                    random_index = random.randint(0, 4)
                    
                    # Retrieve the links at the random indices
                    chosen_link = sliced_links[random_index]
                
                    #^ temporary
                    # print(f"The Unique link (from # urls) for {query} :", chosen_link)
                    
                    return chosen_link
                else:
                    return None
            else:
                print("Failed to retrieve the website content.")
                return None
            
            
            
# Function to retrieve all fragment links
def get_fragment_links(base_url):
    response = requests.get(base_url)
    soup = BeautifulSoup(response.content, 'html.parser')
    fragment_links = []

    # Find all links that point to fragments within the page
    for a in soup.find_all('a', href=True):
        href = a['href']
        if href.startswith('#'):
            fragment_links.append(href)

    return fragment_links



def extract_paragraphs_with_headers(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    
    # Find all paragraphs
    paragraphs_with_headers = []
    
    for paragraph in soup.find_all('p'):
        # Find the closest preceding header (h1 to h6)
        header = paragraph.find_previous_sibling(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        if header:
            header_text = header.get_text(strip=True)
            paragraphs_with_headers.append((header_text, paragraph.get_text(strip=True)))
        else:
            paragraphs_with_headers.append(("", paragraph.get_text(strip=True)))
    
    return "\n".join([f"{header}\n{paragraph}" for header, paragraph in paragraphs_with_headers])



# Collect data from each fragment
def collect_data(base_url, fragment_links):
    all_data = {}
    for fragment in fragment_links:
        full_url = f'{base_url}{fragment}'
        data = extract_paragraphs_with_headers(full_url)
        if data:
            all_data[fragment] = data
        else:
            print(f'Failed to retrieve data from {full_url}')
    return all_data


# Main function
def return_data_withfragments(query):
    if extract_urls(query) is not None:
        url1, url2 = extract_urls(query)
    # Base URL
        bases_urls =  [url1, url2]
        programme_content = []
        for base_url in bases_urls:
            fragment_links = get_fragment_links(base_url)
            if fragment_links:
                all_data = collect_data(base_url, fragment_links)                                       
                for fragment, data in all_data.items():
                    print(f'Data from {fragment}:')
                    print(data)
                    programme_content.append(data)
            else:
                print('No fragment links found.')
        return ("\n".join(programme_content))
    else:
        return None
    
   
# Main function
async def return_data3(query):  
    url1, url2, url3 = await extract3urls(query)  # Await extract3urls here
    
    # Create tasks for fetching content from URLs asynchronously
    if all([url1, url2, url3]):
        task1 = asyncio.create_task(get_url_content(url1))
        task2 = asyncio.create_task(get_url_content(url2))
        task3 = asyncio.create_task(get_url_content(url3))
        
        # Wait for all tasks to complete
        content1, content2, content3 = await asyncio.gather(task1, task2, task3)
            
        # Filter out any None values
        tablo = [content for content in [content1, content2, content3] if content is not None]
        
        if tablo:
            result = '\n\n\n'.join(tablo)
            # print(result)
            # Or return the result if this is inside a function
            return result
    else:
        print("One or more URLs are None")
        return None
        

async def return_data3_orderofresults(query):  
    url1, url2, url3 = await extract3urls_orderofresults(query)  # Await extract3urls here
    
    # Create tasks for fetching content from URLs asynchronously
    if all([url1, url2, url3]):
        task1 = asyncio.create_task(get_url_content(url1))
        task2 = asyncio.create_task(get_url_content(url2))
        task3 = asyncio.create_task(get_url_content(url3))
        
        # Wait for all tasks to complete
        content1, content2, content3 = await asyncio.gather(task1, task2, task3)
            
        # Filter out any None values
        tablo = [content for content in [content1, content2, content3] if content is not None]
        
        if tablo:
            result = '\n\n\n'.join(tablo)
            # print(result)
            # Or return the result if this is inside a function
            return result
    else:
        print("One or more URLs are None")
        return None
        

# Function to load the scraped data from the pickle file
async def load_scraped_data():
    if os.path.exists(PICKLE_FILE):
        with open(PICKLE_FILE, 'rb') as file:
            return pickle.load(file)
    return {}

# Function to save the scraped data to the pickle file
async def save_scraped_data(data):
    with open(PICKLE_FILE, 'wb') as file:
        pickle.dump(data, file)


async def extract_first_url_1(query):
    url = f'https://www.google.com/search?q={requests.utils.quote(query)}'
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }

    async with aiohttp.ClientSession() as session:
        async with session.get(url, headers=headers) as response:
            if response.status == 200:
                html = await response.text()
                soup = BeautifulSoup(html, 'html.parser')
                result = soup.find('div', {'class': 'yuRUbf'})
                if result:
                    print(f"The Unique link for {query} :", result.find('a')['href'])
                    return result.find('a')['href']
                else:
                    print("No results found.")
                    return None
            else:
                print("Failed to retrieve the website content.")
                return None


# Function to scrape content from a URL using the Diffbot API
async def scrape_url(url):
    api_url = f"{DIFFBOT_API_URL}?token={DIFFBOT_API_TOKEN}&url={url}"
    response = requests.get(api_url)
    
    if response.status_code == 200:
        response_data = response.json()
        if 'objects' in response_data and len(response_data['objects']) > 0:
            html_content = response_data['objects'][0].get('html', 'HTML content not available')
            if html_content == 'HTML content not available':
                print(f" for {url} : HTML content not available in the response.")
                return None
            return html_content
        else:
            print(f" for {url} : HTML content not available in the response.")
            return None
    else:
        print(f"Failed to retrieve data from {url}: {response.status_code}")
        return None

# Function to clean HTML content and extract text
async def clean_html(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup.get_text(separator='\n', strip=True)

# Function to get content from a URL, using cached data if available
async def get_url_content(url):
    scraped_data = await load_scraped_data()
    
    if url in scraped_data:
        print(f"Using cached data for {url}")
        return scraped_data[url]
    else:
        content = await scrape_url(url)
        if content:
            text_content = await clean_html(content)
            if text_content:  # Check if the content is not empty
                scraped_data[url] = text_content
                await save_scraped_data(scraped_data)
            return text_content
        return None

async def return_data1(query):
    url1 = await extract_first_url_1(query)
    print(f"for query: {query}  :", url1)
    if url1:
        content = await get_url_content(url1)
        return content
    else:
        print("Nothing found for url1")
        return None

async def return_data1_see(query):
    url1 = await extract_1_url_from_urls(query)
    print(f"for query: {query}  :", url1)
    if url1:
        content = await get_url_content(url1)
        return content
    else:
        print("Nothing found for url1")
        return None



async def retrieve_cgpa_score(content_resume):
    
    # Check if the entity_ruler already exists in the pipeline
    if 'entity_ruler' in nlp.pipe_names:
        ruler = nlp.get_pipe('entity_ruler')
    else:
        ruler = nlp.add_pipe("entity_ruler")

    patterns = [
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.6 CGPA"
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "CGPA 9.6" and "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]},  # Matches "CGPA - 9.26" or "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "9.5 CGPA" or "9.5 CGPA -" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.5 - CGPA" or "9.5 . CGPA" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["-", "–"]}}, {"IS_SPACE": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]}, # Matches "CGPA - 9.26" or "CGPA -9.5" (with optional punctuation and space)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\w*\\d+(\\.\\d+)?"}}, {"LOWER": {"REGEX": "(cgpa|sgpi\\w*)"}}]},
        {"label": "CGPA", "pattern": [{"LOWER": {"REGEX": "(\\w*cgpa|sgpi)"}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?\w*"}}]}
    ]

    ruler.add_patterns(patterns)

    lines = content_resume.split('\n')
    score = "unknown"
    for index, line in enumerate(lines):
        doc = nlp(line)
        found_cgpa = False
        for ent in doc.ents:
            if ent.label_ == 'CGPA':
                found_cgpa = True
                print(ent.text)
                numerical_score = nlp_e(ent.text)
                for token in numerical_score:
                    if token.ent_type_ == 'CARDINAL':
                        score = token.text
                break  # Stop looping over entities once CGPA label is found
        if found_cgpa:
            break  # Stop looping over lines once CGPA label is found

    return score


def retrieve_cgpa_score_sync(content_resume):
    
    # Check if the entity_ruler already exists in the pipeline
    if 'entity_ruler' in nlp.pipe_names:
        ruler = nlp.get_pipe('entity_ruler')
    else:
        ruler = nlp.add_pipe("entity_ruler")

    patterns = [
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.6 CGPA"
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "CGPA 9.6" and "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]},  # Matches "CGPA - 9.26" or "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "9.5 CGPA" or "9.5 CGPA -" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.5 - CGPA" or "9.5 . CGPA" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["-", "–"]}}, {"IS_SPACE": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]}, # Matches "CGPA - 9.26" or "CGPA -9.5" (with optional punctuation and space)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\w*\\d+(\\.\\d+)?"}}, {"LOWER": {"REGEX": "(cgpa|sgpi\\w*)"}}]},
        {"label": "CGPA", "pattern": [{"LOWER": {"REGEX": "(\\w*cgpa|sgpi)"}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?\w*"}}]}
    ]

    ruler.add_patterns(patterns)

    lines = content_resume.split('\n')
    score = "unknown"
    for index, line in enumerate(lines):
        doc = nlp(line)
        found_cgpa = False
        for ent in doc.ents:
            if ent.label_ == 'CGPA':
                found_cgpa = True
                print(ent.text)
                numerical_score = nlp_e(ent.text)
                for token in numerical_score:
                    if token.ent_type_ == 'CARDINAL':
                        score = token.text
                break  # Stop looping over entities once CGPA label is found
        if found_cgpa:
            break  # Stop looping over lines once CGPA label is found

    return score





async def generate_responses_university(res_text,programme, university, international_students, university_description_wikipedia, facilities, research_institutes, ranking, location, culture, professors, practical_learning, fee_structure):
    
    # Get all local variables (including function arguments)
    arguments = locals()
    
    #^ temporary : 
    # for arg_name, arg_value in arguments.items():
    #     if arg_value is not None:
    #         # Print the argument name and its content, with a limit on the content length for readability
    #         print(f"{arg_name}: {' '.join(str(arg_value).split()[:200])}")
    

    completion = ai.ChatCompletion.create(
    #model="gpt-3.5-turbo-16k", 
    model = "gpt-4o-2024-05-13",
    temperature=ai_temp,
    messages = [


{
"role": "user",
"content": f"Imagine you are an Indian student whose resume is {res_text} and you want to study a master programme in {programme} at the university: {university} in Germany. You have to answer from Question 1 to Question 8:"
},
{
"role": "user",
"content": f"Question 1: How do you believe this university's emphasis on practical learning will enhance your academic and professional goals?"
},
{
"role": "user",
"content": f"Your response should be based on {practical_learning} and reflect the style and content of the provided sources. Make sure to include personal reflections and first-person language to make it sound personal and genuine."
},
{
"role": "user",
"content": f"Question 2: How does the expertise and reputation of the faculty at this university influence your decision to attend?"
},
{
"role": "user",
"content": f"Your response should be  based on {university_description_wikipedia} and {professors}"
},
{
"role": "user",
"content": f"Question 3: How do the facilities provided by {university} influence your decision to attend?"
},
{
"role": "user",
"content": f"Your response should be based on {facilities}"
},
{
"role": "user",
"content": f"Question 4: How does the diverse culture at {university} influence your decision to attend the university?"
},
{
"role": "user",
"content": f"Your response should be based on {culture} "
},
{
"role": "user",
"content": f"Question 5: How do the research projects or research centers at {university} influence your decision to attend?"
},
{
"role": "user",
"content": f"Your response should be based on {research_institutes} "
},
{
"role": "user",
"content": f"Question 6: How do the university's ranking and fee structure influence your decision to attend?"
},
{
"role": "user",
"content": f"Your response should be based on {ranking} and {fee_structure}. Please provide the sources. "
},
{
"role": "user",
"content": f"Question 7:  How does the location of this university influence your decision to attend? "
},
{
"role": "user",
"content": f"Your response should be based on {location} and {university_description_wikipedia} "
},
{
"role": "user",
"content": f"Question 1:  How does the size of the student body or the presence of an international student community at this university influence your decision to attend?  "
},
{
"role": "user",
"content": f"Your response should be based on {international_students} "
},
{
"role": "user",
"content": "Provide concrete examples for Questions 1 to 7. For each question, generate 20 lines minimum. I want DETAILED and RELEVANT information that goes beyond a simple sentence."
},

{"role": "user",
"content":"For Questions 1 to 8 The responses should reflect the style and content of the provided sources. Make sure to include personal reflections and first-person language to make it sound personal"},
{
"role": "user",
"content": " For Questions 1 to 8 The responses should be human-like and personal, using first-person language."
},
{"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned', 'intricacies', 'close-knit', 'aligns', 'hands-on', 'enhance', 'foster', 'emphasis', 'boasts'"},
{
"role": "user" , "content" : "I don't want the results in dramatic tone , Instead give results based on the above information i have provided you and try to give results in as simple way as possible "
}
# {
#     "role": "user",
#     "content": "I repeat that the responses should be human-like, meaning very SIMPLE, responses understood by 18-year-old people."
# },






]


    )

    response_out = completion['choices'][0]['message']['content']
    st.write(response_out)
    return response_out
    
    
async def generate_responses_programme(res_text,programme, university, programme_content,  university_no_wikipedia , modules, practical_learning, personal_benefit, professional_growth):


   # Get all local variables (including function arguments)
    arguments = locals()
    
    #^ temporary 
    # for arg_name, arg_value in arguments.items():
    #     if arg_value is not None:
    #         # Print the argument name and its content, with a limit on the content length for readability
    #         print(f"{arg_name}: {' '.join(str(arg_value).split()[:200])}")
    

    completion = ai.ChatCompletion.create(
    #model="gpt-3.5-turbo-16k", 
    model = "gpt-4o-2024-05-13",
    temperature=ai_temp,
    messages = [

{
"role": "user",
"content": f"Imagine you are an Indian student whose resume is {res_text} and you want to study a master programme in {programme} at the university: {university} in Germany. You have to answer from Question 1 to Question 7 :"
},


{
"role": "user",
"content": f"Question 1:  How does the link between the programme {programme_content} and the course curriculum {modules} influence your decision to enroll?  "
},
{
"role": "user",
"content": f"Question 2: How does this program's potential for personal benefit influence your decision to pursue it?  "},
{
"role": "user",
"content": f"Your response should be based on {personal_benefit} "
},
{
"role": "user",
"content": f"Question 3: How does the program's focus on professional growth influence your decision to enroll "},
{
"role": "user",
"content": f"Your response should be based on {professional_growth} "
},

{
"role": "user",
"content": f"Question 4:  How does the emphasis on technical learning and skills in {programme_content} influence your decision to pursue it?  "
},
{
"role": "user",
"content": f"Question 5:  What is the driving motivation behind your choice of this particular program {programme} at {university} ? "},
{
"role": "user",
"content": f"Your response should be based on {university_no_wikipedia} "
},
{
"role": "user",
"content": f"Question 6 : How does your personal experience or interest influence your decision to choose the programme {programme}"
},
{
"role": "user",
"content": f"Your response should be based on your resume {res_text} and {programme_content} "
},
{
"role": "user",
"content": f"Question 7 : How does the emphasis on applying skills in this program influence your decision to enroll? "
},
{
"role": "user",
"content": f"Your response should be based on your resume {res_text} and {programme_content}, and {practical_learning} and {modules} "
},

{
"role": "user",
"content": "Provide concrete examples for Questions 1 to 7. For each question, generate 20 lines minimum. I want DETAILED and RELEVANT information that goes beyond a simple sentence."
},

{"role": "user",
"content":"For Questions 1 to 8 The responses should reflect the style and content of the provided sources. Make sure to include personal reflections and first-person language to make it sound personal"},
{
"role": "user",
"content": " For Questions 1 to 8 The responses should be human-like and personal, using first-person language."
},
{"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned', 'intricacies', 'close-knit', 'aligns', 'hands-on', 'enhance', 'foster', 'emphasis', 'boasts'"},
{
"role": "user" , "content" : "I don't want the results in dramatic tone , Instead give results based on the above information i have provided you and try to give results in as simple way as possible "
}
# {
#     "role": "user",
#     "content": "I repeat that the responses should be human-like, meaning very SIMPLE, responses understood by 18-year-old people."
# },
]


    )

    response_out = completion['choices'][0]['message']['content']
    st.write(response_out)
    return response_out


async def generate_responses_germany(  res_text, programme, university, cooperation_india_germany, quality_education_germany, 
programme_variety_germany, research_opportunities_germany, cultural_experiences_germany, financial_support_germany,
standard_living_germany, location_industries_germany, tuition_fees_germany, postgraduation_opportunities_germany, economic_opportunities_germany):

    
   # Get all local variables (including function arguments)
    arguments = locals()
    
    #^ temporary 
    # for arg_name, arg_value in arguments.items():
    #     if arg_value is not None:
    #         # Print the argument name and its content, with a limit on the content length for readability
    #         print(f"{arg_name}: {' '.join(str(arg_value).split()[:200])}")
    

    completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages = [

        {
        "role": "user",
        "content": f"Imagine you are an Indian student whose resume is {res_text} and you want to study a master programme in {programme} at the university: {university} in Germany."
    },
        
        {"role": "user", "content": f"What factors make Germany an ideal study destination for you? Discuss the following 11 questions with relevant examples, for each question generate 10 lines minimum"},
        {"role": "user", "content": f" Question 1 : Quality Education: How can {quality_education_germany} explain why Germany's high-quality education system is appealing to you."},
        #& Include only the area : 
        {"role": "user", "content": f" Question 2 : Program Variety:  How can {programme_variety_germany} offers the variety of programs available in Germany that aligns with your academic interests and goals in {programme}"},
        #& same here : 
        {"role": "user", "content": f" Question 3 : Economic Opportunities:  How can {economic_opportunities_germany}, Discuss how the opportunities available in Germany align with your career goals in {programme}."},
       
        {"role": "user", "content":  f" Question 4 : Research Opportunities:  Based on {research_opportunities_germany}, Highlight the research opportunities in your field of study that attract you to Germany."},
        # {"role": "user", "content":  f"Question 5 : Employment Opportunities:  Based on {employment_opportunities_germany}, Discuss the employment opportunities available in Germany that align with your career goals."},
        {"role": "user", "content": f"Question 5 :  Based on {cooperation_india_germany}, Mention examples of cooperation between India and Germany relevant to this programme."},
        {"role": "user", "content":  f" Question 6 :  Based on {cultural_experiences_germany}, Cultural Experience: Describe how experiencing German culture and language will enrich your personal and academic growth."},
        { "role": "user", "content":  f"Question 7 : Financial Support:  Based on {financial_support_germany}, Explain how scholarships and financial support options in Germany will help alleviate your financial burden."},
        {"role": "user", "content":   f"Question 8 : Standard of Living:  Based on {standard_living_germany}, Discuss the high standard of living and quality of life that Germany offers to students."},
        {"role": "user", "content": f"Question 9 : Location and Industry Links:  Based on {location_industries_germany}, Explain the strategic advantages of Germany's location and its strong ties to industries relevant to your field."},
        {"role": "user", "content":  f"Question 10 : Tuition Fees:  Based on {tuition_fees_germany}, Discuss the affordability of education in Germany due to its low tuition fees."},
            {"role": "user", "content": f" Question 11 : Post-Graduation Opportunities:  Based on {postgraduation_opportunities_germany}, Highlight the post-graduation opportunities, including visa extensions, that make Germany a favorable choice for your future career prospects."},


    {
        "role": "user",
        "content": "Provide concrete examples for the questions 1 to 11. For each question, generate 10 lines minimum. I want DETAILED and RELEVANT information that goes beyond a simple sentence."
    },

    {"role": "user",
    "content":"For the questions 1 to 11, the responses should reflect the style and content of the provided sources. Make sure to include personal reflections and first-person language to make it sound personal"},

    {
        "role": "user",
        "content": " For the questions 1 to 11, the responses should be human-like and personal, using first-person language."},
    {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned', 'intricacies', 'close-knit', 'aligns', 'hands-on', 'enhance', 'foster', 'emphasis'"},
    {
        "role": "user" , "content" : "I don't want the results in dramatic tone , Instead give results based on the above information i have provided you and try to give results in as simple way as possible "
    }
    # {
    #     "role": "user",
    #     "content": "I repeat that the responses should be human-like, meaning very SIMPLE, responses understood by 18-year-old people."
    # },
    ] )

    response_out = completion['choices'][0]['message']['content']
    st.write(response_out)
    return response_out 
    
       
def create_word_document(phrase, font_name, font_size):
    doc = Document()
    doc.add_paragraph(phrase)

    # Set font properties
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            
        # Set paragraph alignment to justified
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    return doc


def save_doc_to_buffer(doc):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
    
    
def read_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def generate_random_templates(directory_path):
    # List all files in the directory
    files = os.listdir(directory_path)
    # Filter out the PDF files
    pdf_files = [file for file in files if file.endswith('.pdf')]
    
    # Check if there are any PDF files
    if not pdf_files:
        print("No PDF files found in the directory.")
        return
    
    # Randomly select a PDF file
    selected_pdf = random.choice(pdf_files)
    print(f"Selected PDF: {selected_pdf}")
    
    # Full path to the selected PDF file
    pdf_path = os.path.join(directory_path, selected_pdf)
    return read_pdf(pdf_path)
    

def read_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


#! for after if need retrieve contents of articles

def  return_data3_from_diffbot(programme):
    #! {"role": "user", "content": f"Question 13: Talk about cooperations between India and Germany that imagine yourself participating in, in the future. Your response should be based on {cooperation_india_germany}"},
    url = 'https://kg.diffbot.com/kg/v3/dql'
    params = {
        'token': '46db6037f264bbd5e7451b1160d5d88c',
        'query': f'type:Article text:"Cooperation" text:"{programme}" title:"India" title:"Germany"'
    }
    headers = {
        'Accept': 'application/json'
    }

    # Make the GET request
    response = requests.get(url, headers=headers, params=params)

    # Check if the request was successful
    if response.status_code == 200:
        # Parse the JSON response
        data = response.json()
        
        # Access the first result
        if 'data' in data and len(data['data']) > 0:
            first_result = data['data'][0]['entity']
            second_result = data['data'][1]['entity']
            return(first_result['html'], second_result['html'])  # Example of accessing the 'name' field
            
        else:
            return("No results found.")
        
    else:
        # Print the error message if the request failed
        return(f"Request failed with status code {response.status_code}: {response.text}")




#^ THE MAIN PROGRAMME DISPLAYED ON THE STREAMLIT INTERFACE : 

st.markdown("""
# 📝 AI-Powered SOP Generator

Generate a sop letter : JUST READ THE INSTRUCTIONS
"""
)

# radio for upload or copy paste option         
res_format = st.radio(
    "Upload or paste the applicant's resume/key experience",
    ('Upload', 'Paste'))

if res_format == 'Upload':
    # upload_resume
    res_file = st.file_uploader('📁 Upload your resume in pdf or docx format')
    
    if res_file is not None:
        if res_file.name.endswith('.pdf'):
            pdf_reader = PdfReader(res_file)

            # Collect text from pdf
            res_text = ""
            for page in pdf_reader.pages:
                res_text += page.extract_text()
        
        elif res_file.name.endswith('.docx'):
            doc_reader = Document(res_file)

            res_text = ""
            for para in doc_reader.paragraphs:
                res_text += para.text + " "


            for table in doc_reader.tables:
                for row in table.rows:
                    for cell in row.cells:
                        res_text += cell.text + " "
                    res_text += "\n"

        else:
            st.error("Unsupported file format. Please upload a PDF or DOCX file.")
else:
    # Use the pasted contents instead
    res_text = st.text_area('Pasted resume elements')
    
 
 
# st.subheader("If you have a template in mind you can submit it 'Upload'.  Otherwise tick 'Let the software generate the template' and it will be automatically generated ")
# # radio for upload or copy paste option         
# template_format = st.radio(
#     "Do you want to upload or paste the template",
#     ('Upload', 'Paste', 'let the software generate the template'))

# if template_format == 'Upload':     
#         # upload_resume
#     template_file = st.file_uploader('📁 Upload your template in pdf format')
#     if template_file:
#         pdf_reader = PdfReader(template_file)

#         # Collect text from pdf
#         template_text = ""
#         for page in pdf_reader.pages:
#             template_text += page.extract_text()
# elif template_format == 'Paste':
#     # use the pasted contents instead
#     template_text = st.text_input('Pasted template elements')
# else:
#     st.write("Process of selecting a random template")
            
            

with st.form('input_form'):
    # other inputs
    programme = st.text_input('Programme name')
    university = st.text_input('University name')
    ai_temp = st.number_input('AI Temperature (0.0-1.0) Input how creative the API can be',value=.6)
    
    
    
    
    # # submit button
    submitted = st.form_submit_button("Generate the SOP")

# if the form is submitted run the openai completion   
if submitted:    

    #^ generate_responses :                                                                                               
    
    async def main():
        research = university + " " + programme

        

        # Define asynchronous tasks for return_data functions
        
        #! Tasks for University :
        facilities_task = asyncio.create_task(return_data3(f"How do the facilities provided by {university} influence your decision to attend?"))
        research_institutes_task = asyncio.create_task(return_data3_orderofresults(f"How do the research projects or research centers at {university} influence your decision to attend?"))
        university_description_wikipedia_task = asyncio.create_task(return_data1(f"{university} Wikipedia Deutch"))
        professors_task = asyncio.create_task(return_data1(f"good teachers, professors, doctors {university} Research gate"))
        ranking_task = asyncio.create_task(return_data1(f"ranking {university} "))
        location_task = asyncio.create_task(return_data1(f"How does the location of {university} offer something particular and better for students compared to other universities"))
        fee_structure_task = asyncio.create_task(return_data1(f"{university} fee structure")) # here it was return_data3
        international_students_task = asyncio.create_task(return_data3(f"How many international students are at {university}"))
        culture_task = asyncio.create_task(return_data3(f"How does the diverse culture at {university} influence your decision to attend?"))
        
        #* Tasks for the programme : programme_content,  university_no_wikipedia , modules, practical_learning, personal_benefit, professional_growth
        programme_content_task = asyncio.create_task(return_data1(research))
        university_no_wikipedia_task = asyncio.create_task(return_data3_orderofresults(f"{university}"))
        modules_task = asyncio.create_task(return_data3_orderofresults(f"{programme}, {university}, modules"))
        practical_learning_task = asyncio.create_task(return_data3(f"How do you believe the emphasis of {university} on practical learning will enhance your academic and professional goals?"))
        personal_benefit_task = asyncio.create_task(return_data3(f"Why study {programme} in Germany"))
        professional_growth_task = asyncio.create_task(return_data3(f"{programme} professional growth"))
        
        #& Tasks for Germany (retake return_data1 function)
        cooperation_india_germany_task = asyncio.create_task(return_data1_see(f"bilateral cooperation between india and germany, {programme}"))
        quality_education_germany_task = asyncio.create_task(return_data1_see(f"why Germany's high-quality education system is appealing to Indian students"))
        programme_variety_germany_task = asyncio.create_task(return_data1_see(f"Is Germany good for {programme}"))
        research_opportunities_germany_task = asyncio.create_task(return_data1_see(f"Research Opportunities: Highlight the research opportunities in {programme} that attract you to Germany."))
        cultural_experiences_germany_task = asyncio.create_task(return_data1_see(f"Describe how experiencing German culture and language will enrich your personal and academic growth."))
        financial_support_germany_task = asyncio.create_task(return_data1_see(f"Explain how scholarships and financial support options in Germany will help alleviate your financial burden"))
        standard_living_germany_task = asyncio.create_task(return_data1_see(f"Discuss the high standard of living and quality of life that Germany offers to Indian students."))
        location_industries_germany_task = asyncio.create_task(return_data1_see(f"Explain the strategic advantages of Germany's location and its strong ties to industries relevant to {programme}"))
        tuition_fees_germany_task = asyncio.create_task(return_data1_see(f" Discuss the affordability of education in Germany due to its low tuition fees for Indian students."))
        postgraduation_opportunities_germany_task = asyncio.create_task(return_data1_see(f" Highlight the post-graduation opportunities, including visa extensions, that make Germany a favorable choice for your future career prospects in {programme} as an Indian student"))
        economic_opportunities_germany_task = asyncio.create_task(return_data1_see(f"Article about Job Opportunities available in Germany for careers in {programme}"))


    #     #^ Wait for all return_data tasks to complete
       

        (facilities, research_institutes) = await asyncio.gather(facilities_task, research_institutes_task)
        
        await asyncio.sleep(2)
        
        (university_description_wikipedia, professors, cooperation_india_germany) = await asyncio.gather(university_description_wikipedia_task, professors_task,  cooperation_india_germany_task)
        
        await asyncio.sleep(2)
         
        (ranking, location, fee_structure, programme_content, quality_education_germany ) = await asyncio.gather(ranking_task, location_task, fee_structure_task, programme_content_task, quality_education_germany_task)
        
        await asyncio.sleep(2)
        
        ( international_students , modules ) = await asyncio.gather( international_students_task, modules_task)

        await asyncio.sleep(2)
        
        (culture, practical_learning ) = await asyncio.gather( culture_task, practical_learning_task)
        
        await asyncio.sleep(2)
        (  university_no_wikipedia , personal_benefit ) = await asyncio.gather(  university_no_wikipedia_task, personal_benefit_task  )
        
        await asyncio.sleep(2)
        
        ( professional_growth, programme_variety_germany, research_opportunities_germany ) = await asyncio.gather(professional_growth_task , programme_variety_germany_task, research_opportunities_germany_task  )
        
        await asyncio.sleep(2)
        
        ( cultural_experiences_germany, financial_support_germany, standard_living_germany, location_industries_germany, tuition_fees_germany  ) = await asyncio.gather(  cultural_experiences_germany_task, financial_support_germany_task, standard_living_germany_task,  location_industries_germany_task,   tuition_fees_germany_task )
        await asyncio.sleep(2)
        
        ( postgraduation_opportunities_germany , economic_opportunities_germany  ) = await asyncio.gather( postgraduation_opportunities_germany_task,  economic_opportunities_germany_task)
        
        

    
    
    
    
        # Run the response generation functions asynchronously
        response_university = asyncio.create_task(generate_responses_university(
            res_text, programme, university, international_students,
            university_description_wikipedia, facilities,
            research_institutes, ranking, location, culture,
            professors, practical_learning, fee_structure
        ))
        
        response_programme = asyncio.create_task(generate_responses_programme(
             res_text, programme, university,
            programme_content, university_no_wikipedia, modules,
            practical_learning, personal_benefit,
            professional_growth
        ))
        
        
        response_germany = asyncio.create_task(generate_responses_germany(
             res_text, programme, university, cooperation_india_germany, quality_education_germany, programme_variety_germany,
             research_opportunities_germany, cultural_experiences_germany, financial_support_germany, standard_living_germany,
             location_industries_germany,
             tuition_fees_germany, postgraduation_opportunities_germany, economic_opportunities_germany
             
        ))

        # Wait for all response generation tasks to complete
        responses = await asyncio.gather( response_university, response_programme, response_germany)
        
        germany_response = responses[2]
        # # Redirect output to a file
        # with open('output.txt', 'w') as f:
        #     sys.stdout = f
        #     print(type(germany_response))

        # # Reset stdout to default
        # sys.stdout = sys.__stdout__
     
        university_response = responses[0]
        
        programme_response = responses[1]
       
        
        

        
        response_sop = await generate_sop4( generate_random_templates('templates4'), res_text,programme,university,university_response,programme_response,germany_response )

        
        # Combine responses
        response = ''.join(responses) 
        final_sop_and_responses = response + response_sop 
        
        doc_download1 = create_word_document(final_sop_and_responses, 'Arial', 11)
        st.download_button(
            label="Download SOP",
            data=save_doc_to_buffer(doc_download1),
            file_name=f"{res_file.name}_{university}_{programme}_SOP.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    asyncio.run(main())